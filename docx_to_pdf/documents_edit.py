# python_docxの読み込み
from docx import Document

# .docxを読み込んでDocumentオブジェクトの作成
doc = Document("./sample.docx")

# 段落を取得して、その段落に書かれているテキストを抽出・出力
for num, paragraph in enumerate(doc.paragraphs):
    print(num, paragraph.text)

# number = input()
# Y, M, D = input().split()
# address = input()
# group = input()
# address_name = input()

# print(doc.paragraphs)
# dates = doc.paragraphs[2]
# date = dates.text
# date = "令和" + Y + "年" + M + "月" + D + "日"
# dates.text = date
# doc.save("./sample.docx")

