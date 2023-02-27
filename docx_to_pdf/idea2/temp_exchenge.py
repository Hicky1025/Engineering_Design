from docx import Document
import openpyxl as pxl
import re
# from docx2pdf import convert

# 読み込むファイルのリスト
file_list = ["template_1.docx", "template_2.docx", "template_3.xlsx", "template_4.xlsx"]
# インスタンス化したファイルを格納するリスト
file_obj = []
# 共通箇所名を格納するリスト
# 後にdict型で値と関連付ける
commons = []

# 共通箇所を見つける関数
def find_common(obj):

    global commons
    pattern = '\{.*?\}'

    try:
        if bool(obj.tables) == True:
            for table in obj.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for common_text in re.findall(pattern, cell.text):
                            commons.append(common_text)
        
        for paragraph in obj.paragraphs:
            for common_text in re.findall(pattern, paragraph.text):
                commons.append(common_text)

    except AttributeError:
        for sheet_name in obj.sheetnames:
            sheet = obj[sheet_name]
            for row in sheet:
                for cell in row:
                    if type(cell.value) == str:
                        for common_text in re.findall(pattern, cell.value):
                            commons.append(common_text)

for num, file_name in enumerate(file_list):
    # ファイルがdocxのときの処理
    if file_name.endswith(".docx") == True:
        file_obj.append(Document(file_name))
        find_common(file_obj[num])

    # ファイルがxlsxのときの処理
    else:
        file_obj.append(pxl.load_workbook(file_name))
        find_common(file_obj[num])


# 重複した値を除く処理
# set(commons)でもいけるけど、元のリストの順序を保持したいから以下の通りにしてる
# 詳細 : https://note.nkmk.me/python-list-unique-duplicate/
commons = list(dict.fromkeys(commons))

doc_1 = Document("template_1.docx")
doc_2 = Document("template_2.docx")
table = doc_2.tables[0]

xlsx_3 = pxl.load_workbook("./template_3.xlsx")
xlsx_4 = pxl.load_workbook("./template_4.xlsx")
xlsx_value = xlsx_4["Sheet1"]

value = [
    "４産技専管品第２３５号",
    "東京都品川区東大井一丁目１１－７",
    "アイビハイツ南品川管理組合",
    "理事長",
    "高専太郎",
    "令和４年５月１７日",
    "１８６．２０",
    "中央棟４階合同講義室（４３３室）",
    "令和４年５月２８日（土）１０時から１２時まで",
    "５，５０１",
    "U101010",
    "高専品川キャンパス「合同講義室（４階）」の他団体貸付について"
]

# value_list[0] : カラム名, value_list[1] : データ
value_dic = dict(zip(commons, value))

for key in value_dic:
    for paragraph in doc_1.paragraphs:
        paragraph.text = paragraph.text.replace(key, value_dic[key])

    for row in table.rows:
        for cell in row.cells:
            if key in cell.text:
                cell.text = cell.text.replace(key, value_dic[key])

    for row in xlsx_value:
        for cell in row:
            if type(cell.value) == str:
                cell.value = cell.value.replace(key, value_dic[key])

# edit.docx　を新規作成
# doc_1.save("./edit_1.docx")
# doc_2.save("./edit_2.docx")
# xlsx_4.save("./edit_4.xlsx")