from docx import Document
import openpyxl as pxl
import re
import os
# from docx2pdf import convert
# import xlwings as xw
import datetime

# コンテナ上で動いてるOSってホストに依存してるん？それとも
# 保存する時のパス
save_path = os.getcwd() + "/"

# 編集日の日付
edit_day = str(datetime.date.today())

# 読み込むファイルのリスト
file_list = ["template_1.docx", "template_2.docx", "template_3.xlsx", "template_4.xlsx"]

# インスタンス化したファイルを格納するリスト
obj_list = []

# 共通箇所名を格納するリスト
# 後にdict型で値と関連付ける
commons = []

# 共通箇所を見つける関数
def find_common(obj):

    global commons

    # 抽出するパターン
    pattern = '\{.*?\}'

    #　docxの時はtryの処理
    try:
        # docxのオブジェクトがtable（：表）を持てばTrue
        # 以下、tableのセルの値に対して正規表現とマッチする文字列の抽出
        if bool(obj.tables) == True:
            for table in obj.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for common_text in re.findall(pattern, cell.text):
                            commons.append(common_text)
        
        # docxの段落に対して正規表現とマッチする文字列の抽出
        for paragraph in obj.paragraphs:
            for common_text in re.findall(pattern, paragraph.text):
                commons.append(common_text)

    # xlsxはparagraphsパラメータを持たないからAttributeErrorが出る
    # エラーが出るか出ないかでdocxとxlsxを区別してる
    except AttributeError:
        # xlsxの各シートのセルの対して正規表現とマッチする文字列の抽出
        for sheet_name in obj.sheetnames:
            sheet = obj[sheet_name]
            for row in sheet:
                for cell in row:
                    # 文字列型以外をreplaceの引数に与えるとエラーが出る
                    if type(cell.value) == str:
                        for common_text in re.findall(pattern, cell.value):
                            commons.append(common_text)

# ファイルをDocument,workbookオブジェクトに変換する処理
for num, file_name in enumerate(file_list):
    # ファイルがdocxのときの処理
    if file_name.endswith(".docx") == True:
        obj_list.append(Document(save_path + file_name))
        find_common(obj_list[num])

    # ファイルがxlsxのときの処理
    else:
        obj_list.append(pxl.load_workbook(save_path + file_name))
        find_common(obj_list[num])


# 作成したcommonsの重複した値を除く処理
# set(commons)でもいけるけど、元のリストの順序を保持したいから以下の通りにしてる
# 詳細 : https://note.nkmk.me/python-list-unique-duplicate/
commons = list(dict.fromkeys(commons))

# 変数
doc_num = "４産技専管品第２３５号" # {書類番号}
address = "東京都品川区東大井一丁目１１－７" # {宛先住所}
group = "アイビハイツ南品川管理組合" # {組織}
role = "理事長" # {役職}
name = "高専太郎" # {宛名}
app_date = "令和４年５月１７日" # {申請日}
area = "１８６．２０" # {使用面積}
place = "中央棟４階合同講義室（４３３室）" # {使用場所}
limit = "令和４年５月２８日（土）１０時から１２時まで" # {期間}
money = "５，５０１" # {費用}
class_num = "U101010" # {分類記号}
subject = "高専品川キャンパス「合同講義室（４階）」の他団体貸付について" # {文書件名}


# 置換したい値
value = [
    doc_num,
    address,
    group,
    role,
    name,
    app_date,
    area,
    place,
    limit,
    money,
    class_num,
    subject
]

# 共通箇所の名前とデータを辞書型で関連付ける
value_dic = dict(zip(commons, value))

# 置換・保存処理
# 出力はフルパスで
for num, obj_class in enumerate(obj_list):
    for key in value_dic:
        try:
            if bool(obj_class.tables) == True:
                for table in obj_class.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if key in cell.text:
                                cell.text = cell.text.replace(key, value_dic[key])
            
            for paragraph in obj_class.paragraphs:
                paragraph.text = paragraph.text.replace(key, value_dic[key])
            
            save_filename = os.path.splitext(os.path.basename(file_list[num]))[0] + "-" + edit_day + ".docx"
        
        except AttributeError:
            for sheet in obj_class.sheetnames:
                for row in obj_class[sheet]:
                    for cell in row:
                        if type(cell.value) == str:
                            cell.value = cell.value.replace(key, value_dic[key])

            save_filename = os.path.splitext(os.path.basename(file_list[num]))[0] + "-" + edit_day + ".xlsx"
    
    
    if save_filename.endswith(".docx") == True:
        docx_savepath = save_path + save_filename
        pdf_savepath = save_path + save_filename.replace(".docx", ".pdf")
        obj_class.save(docx_savepath)
        # convert(docx_savepath, pdf_savepath)
    
    else:
        xlsx_savepath = save_path + save_filename
        pdf_savepath = save_path + save_filename.replace(".xlsx", ".pdf")
        obj_class.save(xlsx_savepath)
        # workbook = xw.Book(xlsx_savepath)
        # workbook.to_pdf(pdf_savepath)
