# python_doc_1xの読み込み
from docx import document
from docx2pdf import convert

# .doc_1xを読み込んでdoc_1umentオブジェクトの作成
doc_1 = document("./sample.docx")
doc_2 = document("./sample2.docx")

# 変数
# 書類の番号と年月日
number = "４産技専管品第２３５号"
year = "１"
month = "１"
day = "２"

# 依頼者の住所と名前
address_1 = "東京都品川区東大井一丁目１１－７"
address_2 = "アイビハイツ南品川管理組合"
name = "理事長　様"

# 申請年月日
request_year = "５"
request_month = "２"
request_day = "１４"

# 使用の詳細
place_name = "東京都立産業技術高等専門学校高専品川キャンパス"
place_address = "東京都品川区東大井１丁目１０番４０号"
place_type = "建物（会議室部分）"
place_item = "事務所建"
place_area = "１８６．２０"
place_part = "中央棟４階合同講義室（４３３室）"

# 期限
limit_year = "４"
limit_month = "５"
limit_day = "２８"
limit_dow = "土"
limit_time_1 = "１０"
limit_time_2 = "１２"

# 以下、編集処理
# 上の変数の値を適切な箇所に代入
doc_1_number = doc_1.paragraphs[1]
doc_1_number.text = number

dates = doc_1.paragraphs[2]
dates.text = "令和" + year + "年" + month + "月" + day + "日"

client_address_1 = doc_1.paragraphs[6]
client_address_1.text = "　" + address_1
client_address_2 = doc_1.paragraphs[7]
client_address_2.text = "　" + address_2
client_name = doc_1.paragraphs[8]
client_name.text = "　　　　" + name

request_date = doc_1.paragraphs[10]
request_date.text = "令和" + request_year +  "年" + request_month + "月" + request_day + "日付で申請のありました東京都公立大学法人資産の使用については、下記のとおり、許可します。"

host_place_name = doc_1.paragraphs[19]
host_place_name.text = "　　　　名　　称　　" + place_name
host_place_address = doc_1.paragraphs[21]
host_place_address.text = "　　　　所　　在　　" + place_address
host_place_type = doc_1.paragraphs[22]
host_place_type.text = "　　　　種　　類　　" + place_type
host_place_item = doc_1.paragraphs[23]
host_place_item.text = "　　　　種　　目　　" + place_item
host_place_area = doc_1.paragraphs[24]
host_place_area.text = "　　　　使用面積　　" + place_area + "㎡"
host_place_part = doc_1.paragraphs[25]
host_place_part.text = "　　　　使用部分　　" + place_part

limit_date = doc_1.paragraphs[28]
limit_date.text = "第２　　使用期間は、令和" + limit_year + "年" + limit_month + "月" + limit_day + "日（" + limit_dow + "）" + limit_time_1 + "時から" + limit_time_2 + "時までとする。"

# 編集したwardファイルの保存
doc_1.save("./sample.doc_1x")

# wardファイルのPDF化
inputfile = "./sample.doc_1x"
outputfile = "./sample.pdf"
convert(inputfile, outputfile)