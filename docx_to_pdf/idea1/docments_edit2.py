# python_docxの読み込み
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx2pdf import convert

# .docxを読み込んでDocumentオブジェクトの作成
doc = Document("./sample2.docx")
for i in doc.element.body.iter():
    print(i)
    # for j in i:
    #     print(j.text)

# tbl = doc.tables[0]
# for row in tbl.rows:
#     values = []
#     for cell in row.cells:
#         values.append(cell.text)
#     print(values)
# cl = tbl.cell(0, 1)
# for i in cl.paragraphs:
#     i.text = "保存期間"
#     i.alignment = WD_ALIGN_PARAGRAPH.CENTER


# doc.save("./sample2.docx")